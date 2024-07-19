import docx
from docx.oxml.shared import OxmlElement
from tqdm import tqdm

from psalm_scraper import grab_psalm


def create_psalms_booklet(template_path, output_path):
    # Create a new document based on the template
    doc = docx.Document(template_path)

    start_indented = False

    for psalm_number in tqdm(list(range(1, 151))):
        # Grab psalm online
        psalm = grab_psalm(psalm_number)

        # Add title
        p = doc.add_paragraph(style='Psalm Title')
        p.add_run(f"Psalm {psalm['number']} ")

        # Add purpose
        if "purpose" in psalm:
            purpose_run = p.add_run(f"{psalm['purpose']}")
            purpose_run.style = 'Psalm Purpose'
            purpose_run.bold = False

        # Add verses
        for i_verse, (verse_number, verse_rows) in enumerate(psalm['verses'].items()):
            style = "Verse Non-indented" if i_verse % 2 == int(start_indented) else "Verse Indented"
            p = doc.add_paragraph(style=style)
            p.add_run(f"{verse_number}").font.superscript = True
            for i_row, verse_row in enumerate(verse_rows):
                if i_row == 0:
                    run = p.add_run("\t" + verse_row)
                else:
                    run._element.append(OxmlElement('w:br'))
                    run = p.add_run(verse_row)
        start_indented = style == "Verse Non-indented"  # If last verse was not indented, start next psalm indented

        doc.add_paragraph()

    doc.save(output_path)


create_psalms_booklet('psalms_template.docx', 'psalms_booklet.docx')
